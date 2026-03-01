from docx.api import Document
from typing import Any
import logging

logger = logging.getLogger(__name__)


def translate_docx(input_path: str, output_path: str, translator: Any, target_lang: str):
    """Translate a DOCX file using batch translation while preserving
    run-level formatting (bold, italic, font, size, colour, etc.)."""
    doc = Document(input_path)

    # --- collect every run that has translatable text ---
    RunRef = tuple  # (run_object,)
    refs: list[Any] = []
    texts: list[str] = []

    for para in doc.paragraphs:
        for run in para.runs:
            if run.text and run.text.strip():
                refs.append(run)
                texts.append(run.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text and run.text.strip():
                            refs.append(run)
                            texts.append(run.text)

    # --- batch-translate all collected texts ---
    if texts:
        try:
            translated = translator.translate_batch(texts, target_lang)
        except Exception:
            logger.exception("Batch translation failed for DOCX; falling back to per-item")
            translated = []
            for t in texts:
                try:
                    r = translator.translate_text(t, target_lang)
                    translated.append(r if r is not None else t)
                except Exception:
                    logger.exception("Per-item fallback also failed")
                    translated.append(t)

        # --- write results back into the runs ---
        errors = 0
        for run, orig, tr in zip(refs, texts, translated):
            if tr is None:
                tr = orig
                errors += 1
            run.text = tr
    else:
        errors = 0

    doc.save(output_path)
    if errors:
        raise RuntimeError(f"Translation completed with {errors} failed spans")
